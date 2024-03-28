# -*- coding: utf-8 -*-
import json
import openai 
import time
import sys
import docx

RED = "\033[0;31m"
GREEN = "\033[0;32m"
YELLOW = "\033[0;33m"
  
def load_api_key(secrets_file="secrectsChatGPT.json"):
    with open(secrets_file) as f:
        secrets = json.load(f)
    return secrets["OPENAI_API_KEY"]

api_key = load_api_key()
openai.api_key = api_key

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # this is the degree of randomness of the model's output
    )
    return response.choices[0].message["content"]

#interviews = getText("C:\\Users\\sandr\\Desktop\\HiddenProfiles\\NoCriticalInfo-A.docx")

def prompt(interviews):
    return f"""
        Tu és um investigador criminal e queres listar todas as pistas necessárias para solucionar um crime de homicídio usando as entrevistas já realizadas, delimitadas ````.
        Para cada pista escreva uma frase curta com o maximo de 10 palavras.
        Interviews: ```{interviews}```
        """



doc = docx.Document("C:\\Users\\sandr\\Desktop\\HiddenProfiles\\NoCriticalInfo-A.docx")
fullText = []
temp = ""
for para in doc.paragraphs:
    temp += para.text
    if "batata" in para.text:
        temp = temp.replace('batata', '')
        fullText += [temp]
        temp = ""


for t in fullText:
    response = get_completion(prompt(t))
    print(response)
    print("\n")